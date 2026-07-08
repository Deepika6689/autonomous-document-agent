"""
Builds a polished .docx from the agent's plan + drafted sections, using
python-docx. Kept dependency-light (no docx-js/Node) since this whole
project is a standalone Python service.
"""

import os
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)


def _add_title_page(doc: Document, title: str, subtitle: str):
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(subtitle)
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(date.today().strftime("%B %d, %Y"))
    date_run.font.size = Pt(11)
    date_run.font.italic = True

    doc.add_paragraph()


def _add_assumptions_box(doc: Document, assumptions: list):
    if not assumptions:
        return
    heading = doc.add_paragraph()
    run = heading.add_run("Assumptions Made by the Agent")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT_COLOR
    for a in assumptions:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(a)
    doc.add_paragraph()


def _add_table(doc: Document, table_spec: dict):
    headers = table_spec.get("headers") or []
    rows = table_spec.get("rows") or []
    if not headers or not rows:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if i < len(cells):
                cells[i].text = str(val)
    doc.add_paragraph()


def build_docx(document_type: str, user_request: str, assumptions: list,
               sections: dict, output_path: str) -> str:
    doc = Document()

    # Base document margins/font (kept simple + portable, no template dependency)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_title_page(doc, document_type, f"Auto-generated in response to: \"{user_request.strip()}\"")
    _add_assumptions_box(doc, assumptions)

    for section_name, content in sections.items():
        h = doc.add_heading(section_name, level=1)
        for run in h.runs:
            run.font.color.rgb = ACCENT_COLOR

        for para in content.get("paragraphs", []) or []:
            doc.add_paragraph(para)

        for bullet in content.get("bullets", []) or []:
            doc.add_paragraph(bullet, style="List Bullet")

        table_spec = content.get("table")
        if table_spec:
            _add_table(doc, table_spec)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
